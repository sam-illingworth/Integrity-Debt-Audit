import streamlit as st
import google.generativeai as genai
import json
from pypdf import PdfReader
from docx import Document
from fpdf import FPDF
import io
import requests
from bs4 import BeautifulSoup
import re
from google.api_core import exceptions

# 1. Configuration and Mobile CSS
st.set_page_config(page_title="Integrity Debt Diagnostic", page_icon="⚖️", layout="wide")

st.markdown("""
    <style>
    .stApp { 
        background-color: white; 
        color: black; 
    }
    header {visibility: hidden;}
    .reportview-container { background: white; }
    p, span, h1, h2, h3, h4, li { color: black !important; }
    
    /* Fix all input elements */
    .stTextInput > div > div > input {
        background-color: white !important;
        border: 2px solid #d0d0d0 !important;
        color: black !important;
    }
    
    /* Fix select/dropdown boxes */
    .stSelectbox > div > div > div {
        background-color: white !important;
        color: black !important;
    }
    
    .stSelectbox [data-baseweb="select"] {
        background-color: white !important;
    }
    
    .stSelectbox [data-baseweb="select"] > div {
        background-color: white !important;
        border: 2px solid #d0d0d0 !important;
    }
    
    /* Fix file uploader */
    [data-testid="stFileUploader"] {
        background-color: white !important;
        border: 2px dashed #d0d0d0 !important;
    }
    
    [data-testid="stFileUploader"] section {
        background-color: white !important;
        border: none !important;
    }
    
    /* Fix text area */
    .stTextArea textarea {
        background-color: white !important;
        border: 2px solid #d0d0d0 !important;
        color: black !important;
    }
    
    /* Fix buttons - simple white background with dark text */
    .stButton > button {
        background-color: white !important;
        color: #212529 !important;
        border: 2px solid #212529 !important;
        padding: 12px 24px !important;
        font-weight: 600 !important;
    }
    
    .stButton > button:hover {
        background-color: #f8f9fa !important;
        border: 2px solid #000000 !important;
    }
    
    /* Fix form submit button - same simple style */
    .stFormSubmitButton > button {
        background-color: white !important;
        color: #212529 !important;
        border: 2px solid #212529 !important;
        padding: 12px 24px !important;
        font-weight: 600 !important;
    }
    
    .stFormSubmitButton > button:hover {
        background-color: #f8f9fa !important;
        border: 2px solid #000000 !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ==============================================================================
# DATA STRUCTURES FOR ENHANCED PDF GENERATION
# ==============================================================================

PEDAGOGICAL_CONTEXT = {
    'Final product weighting': 'When all the marks sit on a single final deadline, students feel pressure to deliver a polished product at any cost. That makes AI tempting. If you spread the marks across drafts, feedback responses, and planning stages, you reward the actual learning journey. Students cannot fake sustained engagement over weeks. The process becomes more valuable than the final polish.',
    'Iterative documentation': 'AI tools hide their tracks. They produce seamless, polished text instantly. Real human learning is messy. It involves false starts, abandoned ideas, and gradual improvements. When you ask students to show this messiness through lab books, draft annotations, or revision logs, you make AI automation much harder. The struggle to develop an idea cannot be faked overnight.',
    'Contextual specificity': 'AI models train on generic textbook knowledge. They excel at broad theoretical questions. They struggle with specific contexts like your classroom debate last Tuesday, the local council decision students witnessed, or the guest speaker who challenged conventional thinking. When you anchor assessments in unique moments and places, AI cannot replicate that specificity without the student having actually been there.',
    'Reflective criticality': 'Generic professional reflection is easy to automate. AI can produce convincing statements like \'I learned the importance of teamwork.\' But genuine reflection requires vulnerability. It asks students to describe specific moments of confusion, physical sensations of discomfort, and how their values shifted. AI cannot fabricate lived experience. Only humans can connect theory to what it felt like to fail and try again.',
    'Temporal friction': 'If an assessment can be completed in one night, it will be. AI thrives on speed. Building in time delays makes automation much harder. Longitudinal studies require data collection over weeks. Peer review cycles force students to wait for feedback before progressing. Sequential deadlines prevent sprinting. When time itself becomes part of the assessment structure, students must engage with the material over the long term.',
    'Multimodal evidence': 'Text is AI\'s native format. Word documents are trivially easy to automate. But when you ask for a hand-drawn concept map photographed and annotated, a 3-minute audio reflection, or a physical model presented in class, you move students beyond the textbox. These modes require different types of engagement. They add layers of human authenticity that pure text cannot match.',
    'Explicit AI interrogation': 'Banning AI does not work. Students use it anyway. Instead, bring the tool into the classroom as something to study and critique. Ask students to generate AI content, then spend the assessment identifying its errors, biases, and missing nuance. This transforms AI from a hidden shortcut into the object of critical analysis. It teaches students that human expertise matters because AI fails in predictable ways.',
    'Real-time defence': 'Live conversation reveals understanding in ways written text cannot. AI can draft perfect scripts. It cannot handle spontaneous questions about methodology. It cannot justify choices on the spot. A 10-minute viva, a live presentation with Q&A, or peer critique sessions force students to think aloud. This shifts the dynamic from surveillance to dialogue. Authentic ownership becomes visible through unscripted speech.',
    'Social and collaborative labour': 'Automation is solitary. Genuine learning thrives in groups. When students must explain their thinking to peers, respond to challenges, give feedback, and negotiate ideas together, they create witnesses to their process. Verified group work, observed collaboration, and graded peer feedback make it much harder to outsource the task to AI. Social friction creates accountability.',
    'Data recency': 'AI training data is always historical. It knows the past well but struggles with the present. When you ask students to analyse this morning\'s headlines, this week\'s policy change, or datasets released in the last fortnight, you create a knowledge barrier. AI will hallucinate or guess. Students must engage with current events. This tethers assessments to the living world rather than static textbook content.'
}

IMPROVEMENT_ACTIONS = {
    'Final product weighting': [
        'Allocate 20% of marks to the initial research plan',
        'Require a \'response to feedback\' log as part of the final submission',
        'Use scaffolded deadlines throughout the module'
    ],
    'Iterative documentation': [
        'Mandate the use of a weekly digital or physical lab book/process log',
        'Include a \'failed paths\' section where students explain ideas they abandoned',
        'Encourage version control or tracked changes as evidence'
    ],
    'Contextual specificity': [
        'Reference a specific guest speaker or seminar debate in the prompt',
        'Require students to apply theory to a local community issue',
        'Update prompts every semester to reflect the current political or social climate'
    ],
    'Reflective criticality': [
        'Ask for \'I\' statements and specific sensory details of the learning experience',
        'Encourage non-standard formats like reflective poetry or audio diaries',
        'Require students to link specific personal values to the academic content'
    ],
    'Temporal friction': [
        'Build in a mandated peer-review cycle in week 6 of a 12-week module',
        'Require data collection that occurs at specific intervals',
        'Design tasks that require sequential steps that cannot be bypassed'
    ],
    'Multimodal evidence': [
        'Replace one essay with a 5-minute narrated video or podcast',
        'Require hand-drawn diagrams or mind maps to be scanned and included',
        'Use pitch sessions where students explain concepts verbally'
    ],
    'Explicit AI interrogation': [
        'Set an assessment where the goal is to break the AI\'s logic',
        'Task students with fact-checking a synthetic essay',
        'Discuss the ethical and environmental costs of AI in the classroom'
    ],
    'Real-time defence': [
        'Implement 10-minute \'flash vivas\' for high-stakes work',
        'Use in-class critique sessions where peers question each other\'s methodology',
        'Record short verbal feedback loops between tutor and student'
    ],
    'Social and collaborative labour': [
        'Grade the quality of the feedback a student gives to their teammates',
        'Use collaborative drafting sessions during seminar time',
        'Require a reflective log on the challenges of the group dynamic'
    ],
    'Data recency': [
        'Use \'this morning\'s headlines\' as the basis for a theory application',
        'Require students to use the most recent 6 months of a specific journal',
        'Set tasks based on live, streaming data or current social media trends'
    ]
}


# 2. Professional PDF Class - UPDATED FOR ORPHAN PREVENTION
class IntegrityPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(20, 20, 20)
        self.set_auto_page_break(auto=True, margin=20)
        self.primary_color = (45, 77, 74)     # #2D4D4A
        self.bg_cream = (247, 243, 233)      # #F7F3E9
        self.text_color_val = (45, 77, 74)   # #2D4D4A
        self.accent_blue = (52, 152, 219)    
        self.success = (39, 174, 96)         
        self.warning = (243, 156, 18)        
        self.danger = (231, 76, 60)          
        self.light_grey = (245, 247, 250)    

    def header(self):
        self.set_fill_color(*self.bg_cream)
        self.rect(0, 0, 210, 35, 'F')
        self.set_y(10)
        self.set_font('helvetica', 'B', 18)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, 'Integrity Debt Audit Report', 0, 1, 'L')
        self.set_font('helvetica', 'I', 11)
        self.cell(0, 8, 'Framework by Professor Sam Illingworth', 0, 1, 'L')
        self.ln(20)

    def footer(self):
        self.set_y(-18)
        self.set_draw_color(200, 200, 200)
        self.line(20, self.get_y(), 190, self.get_y())
        self.set_font('helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.set_x(20)
        self.cell(0, 10, f'Integrity Debt Audit Report | Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    def safe_text(self, text):
        if not text: return "N/A"
        text_str = str(text)
        replacements = {'\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"', '\u2026': '...', '\u2022': '*'}
        for char, replacement in replacements.items():
            text_str = text_str.replace(char, replacement)
        return text_str.encode('latin-1', 'replace').decode('latin-1')

    def check_page_break(self, height_needed):
        if self.get_y() + height_needed > self.page_break_trigger:
            self.add_page()

    def add_summary(self, actual, score, improvements, doc_context):
        self.set_x(20)
        self.set_font('helvetica', 'B', 16)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, "Executive Summary", 0, 1)
        self.ln(2)
        self.set_fill_color(*self.light_grey)
        self.set_draw_color(220, 220, 220)
        box_y = self.get_y()
        self.rect(20, box_y, 170, 45, 'FD')
        self.set_xy(25, box_y + 5)
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(35, 8, "Context:", 0, 0)
        self.set_font('helvetica', '', 11)
        self.set_text_color(*self.text_color_val)
        self.multi_cell(125, 8, self.safe_text(doc_context))
        self.set_xy(25, box_y + 28)
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(35, 8, "Integrity Score:", 0, 0)
        # HIGH score (40-50) = GOOD (green), LOW (10-20) = BAD (red)
        sc_col = self.success if score >= 40 else self.warning if score >= 25 else self.danger
        self.set_font('helvetica', 'B', 14)
        self.set_text_color(*sc_col)
        self.cell(30, 8, f"{score}/50", 0, 0)
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(35, 8, "Susceptibility:", 0, 0)
        self.set_font('helvetica', 'B', 12)
        self.set_text_color(*sc_col)
        self.cell(30, 8, actual, 0, 1)
        self.set_y(box_y + 50)
        self.set_x(20)
        self.ln(5)
        self.set_font('helvetica', 'B', 14)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, "Top 3 Priority Improvements", 0, 1)
        self.set_font('helvetica', '', 11)
        for i, imp in enumerate(improvements[:3], 1):
            self.set_x(20)
            self.set_text_color(*self.accent_blue)
            self.cell(10, 8, f"{i}.", 0, 0)
            self.set_text_color(*self.text_color_val)
            self.multi_cell(0, 8, self.safe_text(imp))
        self.ln(10)
    
    def add_score_interpretation(self, total_score):
        """Visual scoring scale and interpretation - always on new page"""
        # Always start on fresh page
        self.add_page()
        
        self.set_font('helvetica', 'B', 14)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, 'Understanding Your Score', 0, 1)
        self.ln(3)
        
        # Visual scale using cells with fills (more reliable than rectangles)
        scale_y = self.get_y()
        
        # Three zones: 10-24 (red), 25-39 (amber), 40-50 (green)
        zone1_width = 62
        zone2_width = 62
        zone3_width = 46
        
        # Simpler approach: don't fill, just draw bordered cells with text
        # Then draw colored rectangles BEHIND them
        self.set_xy(20, scale_y)
        
        # First draw the colored backgrounds
        self.set_fill_color(231, 76, 60)  # Red
        self.rect(20, scale_y, zone1_width, 12, 'F')
        
        self.set_fill_color(243, 156, 18)  # Amber
        self.rect(20 + zone1_width, scale_y, zone2_width, 12, 'F')
        
        self.set_fill_color(39, 174, 96)  # Green
        self.rect(20 + zone1_width + zone2_width, scale_y, zone3_width, 12, 'F')
        
        # Now draw borders
        self.set_draw_color(255, 255, 255)  # White borders for contrast
        self.set_line_width(1)
        self.rect(20, scale_y, zone1_width, 12, 'D')
        self.rect(20 + zone1_width, scale_y, zone2_width, 12, 'D')
        self.rect(20 + zone1_width + zone2_width, scale_y, zone3_width, 12, 'D')
        
        # Now overlay text on top
        self.set_xy(20, scale_y + 2)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(255, 255, 255)
        self.cell(zone1_width, 8, '10-24: Critical', 0, 0, 'C')
        
        self.set_xy(20 + zone1_width, scale_y + 2)
        self.cell(zone2_width, 8, '25-39: Moderate', 0, 0, 'C')
        
        self.set_xy(20 + zone1_width + zone2_width, scale_y + 2)
        self.cell(zone3_width, 8, '40-50: Resilient', 0, 0, 'C')
        
        # Mark their score with bigger, clearer text ABOVE the scale
        if total_score >= 10:
            # Calculate position based on score
            total_scale_width = 170
            score_position = 20 + ((total_score - 10) / 40) * total_scale_width
            arrow_y = scale_y - 12
            
            # Draw arrow
            self.set_draw_color(0, 0, 0)
            self.set_line_width(1)
            self.line(score_position, arrow_y + 8, score_position, scale_y)
            
            # Score label
            self.set_xy(score_position - 18, arrow_y - 8)
            self.set_font('helvetica', 'B', 12)
            self.set_text_color(0, 0, 0)
            self.cell(36, 6, f'Your score: {total_score}', 0, 0, 'C')
        
        # NOW move cursor to AFTER the scale boxes and add lots of space
        self.set_y(scale_y + 12)
        self.ln(12)
        
        # Interpretation box
        if total_score >= 40:
            title = 'Strong Position: Keep Building'
            message = 'Your assessment shows good resistance to AI automation. Focus on your few remaining vulnerabilities and share your practices with colleagues.'
            box_color = self.success
        elif total_score >= 25:
            title = 'Moderate Risk: Redesign Recommended'
            message = 'Your assessment has protective features but significant gaps remain. Students can partially automate your work. Prioritise the improvements below, starting with your lowest-scoring categories.'
            box_color = self.warning
        else:
            title = 'High Risk: Urgent Action Needed'
            message = 'Your assessment can be largely automated by AI with minimal human input. This threatens the value of your qualification. Begin redesigning immediately, focusing on your three weakest categories.'
            box_color = self.danger
        
        self.set_fill_color(*box_color)
        self.set_text_color(255, 255, 255)
        self.set_font('helvetica', 'B', 11)
        self.cell(0, 8, f' {title}', 0, 1, 'L', True)
        
        self.set_fill_color(250, 250, 250)
        self.set_text_color(*self.text_color_val)
        self.set_font('helvetica', '', 10)
        self.multi_cell(0, 6, message, 1, 'L', True)
        self.ln(8)

    def add_category(self, name, score, critique, question, quote):
        # Anchor check to prevent orphaned headlines
        self.check_page_break(60) 
        # HIGH scores (4-5) are GOOD (green), LOW (1-2) are BAD (red)
        accent = self.success if score >= 4 else self.warning if score == 3 else self.danger
        status = "RESILIENT" if score >= 4 else "MODERATE" if score == 3 else "VULNERABLE"
        self.set_x(20)
        start_y = self.get_y()
        self.set_fill_color(*accent)
        self.rect(20, start_y + 2, 2, 8, 'F')
        self.set_xy(23, start_y)
        self.set_font('helvetica', 'B', 12)
        self.set_text_color(*self.primary_color)
        self.cell(120, 12, f" {self.safe_text(name)}", 0, 0, 'L')
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*accent)
        self.cell(0, 12, f"Score: {score}/5 | {status}", 0, 1, 'R')
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, "Critique:", 0, 1)
        self.set_x(20)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        self.multi_cell(0, 6, self.safe_text(critique))
        self.ln(3)
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, "Dialogue Question:", 0, 1)
        self.set_x(20)
        self.set_font('helvetica', 'I', 10)
        self.set_text_color(*self.text_color_val)
        self.multi_cell(0, 6, self.safe_text(question))
        self.ln(3)
        self.set_x(20)
        self.set_font('helvetica', 'B', 9)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, "Evidence Reference:", 0, 1)
        self.set_x(20)
        self.set_font('courier', '', 9) 
        self.set_text_color(80, 80, 80)
        self.set_fill_color(250, 250, 250)
        self.set_draw_color(230, 230, 230)
        self.multi_cell(0, 5, f"\"{self.safe_text(quote)}\"", 1, 'L', True)
        self.ln(8)

    def add_contact_box(self):
        """Contact box with no duplication - keep together on page"""
        # If not enough room for entire contact section, start new page
        if self.get_y() + 50 > self.page_break_trigger:
            self.add_page()
        self.ln(10)
        self.set_x(20)
        self.set_fill_color(*self.bg_cream)
        self.set_text_color(*self.primary_color)
        self.set_font('helvetica', 'B', 12)
        self.cell(0, 10, ' Need Support with Implementation?', 0, 1, 'L', True)
        self.set_fill_color(245, 247, 250)
        self.set_text_color(*self.text_color_val)
        self.set_font('helvetica', '', 10)
        self.set_x(20)
        # Main text in bordered box
        self.set_fill_color(245, 247, 250)
        self.set_draw_color(220, 220, 220)
        contact_text = "As a Full Professor with over 20 years experience working in higher education, I can help you interpret your results and redesign your assessments for AI resilience."
        self.multi_cell(0, 6, contact_text, 1, 'L', True)
        self.ln(3)
        
        # Strategy call - label and link on same line
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.text_color_val)
        self.cell(40, 6, 'Book a strategy call:', 0, 0)
        self.set_font('helvetica', 'U', 10)
        self.set_text_color(0, 0, 255)
        self.cell(0, 6, 'sam.illingworth@gmail.com', 0, 1)
        
        # Community - label and link on same line
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.text_color_val)
        self.cell(40, 6, 'Join Slow AI:', 0, 0)
        self.set_font('helvetica', 'U', 10)
        self.set_text_color(0, 0, 255)
        self.cell(0, 6, 'theslowai.substack.com', 0, 1)
        
        self.ln(5)

    def add_category_deep_dive(self, name, score, critique, question, quote, pedagogical_context, actions):
        """Enhanced category with pedagogical context and actions - keeps section together on page"""
        # Calculate approximate height needed for entire section
        context_lines = len(pedagogical_context) // 90  # Approx chars per line
        actions_lines = len(actions) * 2  # Each action ~2 lines
        total_height = 60 + (context_lines * 6) + (actions_lines * 6)  # Base + context + actions
        
        # If not enough room, start new page to keep section together
        if self.get_y() + total_height > self.page_break_trigger:
            self.add_page()
        
        accent = self.success if score >= 4 else self.warning if score == 3 else self.danger
        status = 'RESILIENT' if score >= 4 else 'MODERATE' if score == 3 else 'VULNERABLE'
        
        # Header with colored bar
        self.set_x(20)
        start_y = self.get_y()
        self.set_fill_color(*accent)
        self.rect(20, start_y + 2, 2, 8, 'F')
        self.set_xy(23, start_y)
        self.set_font('helvetica', 'B', 12)
        self.set_text_color(*self.primary_color)
        self.cell(120, 12, f' {self.safe_text(name)}', 0, 0, 'L')
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*accent)
        self.cell(0, 12, f'Score: {score}/5 | {status}', 0, 1, 'R')
        
        # AI-generated critique of THEIR assessment
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, 'Your Assessment Analysis:', 0, 1)
        self.set_x(20)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        self.multi_cell(0, 6, self.safe_text(critique))
        self.ln(3)
        
        # Pedagogical theory connected to their score
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        
        if score >= 4:
            theory_header = 'Why This Strength Matters:'
        elif score == 3:
            theory_header = 'Why Addressing This Gap Matters:'
        else:
            theory_header = 'Why This Vulnerability Is Critical:'
        
        self.cell(0, 6, theory_header, 0, 1)
        self.set_x(20)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        self.set_fill_color(*self.bg_cream)
        self.multi_cell(0, 6, self.safe_text(pedagogical_context), 0, 'L', True)
        self.ln(3)
        
        # Specific improvement actions
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, 'Practical Steps to Strengthen This:', 0, 1)
        for i, action in enumerate(actions, 1):
            self.set_x(20)
            self.set_font('helvetica', '', 10)
            self.set_text_color(*self.accent_blue)
            self.cell(7, 6, f'{i}.', 0, 0)
            self.set_text_color(*self.text_color_val)
            self.multi_cell(0, 6, self.safe_text(action))
        self.ln(2)
        
        # Reflection question for curriculum conversations
        self.set_x(20)
        self.set_font('helvetica', 'B', 10)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, 'Discuss With Your Team:', 0, 1)
        self.set_x(20)
        self.set_font('helvetica', 'I', 10)
        self.set_text_color(*self.text_color_val)
        self.multi_cell(0, 6, self.safe_text(question))
        self.ln(3)
        
        # Evidence quote from their assessment
        self.set_x(20)
        self.set_font('helvetica', 'B', 9)
        self.set_text_color(*self.primary_color)
        self.cell(0, 6, 'Where We Found This in Your Brief:', 0, 1)
        self.set_x(20)
        self.set_font('courier', '', 9)
        self.set_text_color(80, 80, 80)
        self.set_fill_color(250, 250, 250)
        self.set_draw_color(230, 230, 230)
        self.multi_cell(0, 5, f'"{self.safe_text(quote)}"', 1, 'L', True)
        self.ln(8)

    def add_action_plan(self, results):
        """Pre-filled action plan for 3 weakest categories"""
        # Always start action plan on new page
        self.add_page()
        self.set_font('helvetica', 'B', 16)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, 'Priority Action Plan', 0, 1)
        self.ln(3)
        self.set_font('helvetica', '', 11)
        self.multi_cell(0, 6, 'Focus on your three weakest categories:')
        self.ln(5)
        
        sorted_cats = sorted(results.items(), key=lambda x: x[1]['verified_score'])
        for rank, (cat_name, data) in enumerate(sorted_cats[:3], 1):
            self.set_font('helvetica', 'B', 12)
            self.set_text_color(*self.primary_color)
            self.cell(0, 8, f'Priority {rank}: {cat_name} ({data["verified_score"]}/5)', 0, 1)
            actions = IMPROVEMENT_ACTIONS.get(cat_name, [])
            for i, action in enumerate(actions, 1):
                self.set_x(25)
                self.set_font('helvetica', 'B', 10)
                self.set_text_color(*self.accent_blue)
                self.cell(7, 6, f'{i}.', 0, 0)
                self.set_font('helvetica', '', 10)
                self.set_text_color(*self.text_color_val)
                self.multi_cell(0, 6, self.safe_text(action))
            self.ln(5)


    
    def add_next_steps(self):
        """Practical guidance on what to do with this report"""
        self.check_page_break(70)
        self.add_page()
        self.set_font('helvetica', 'B', 16)
        self.set_text_color(*self.primary_color)
        self.cell(0, 10, 'What To Do Next', 0, 1)
        self.ln(3)
        
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        intro = "You now have a detailed diagnostic of your assessment. Here is how to use this report effectively:"
        self.multi_cell(0, 6, intro)
        self.ln(5)
        
        # Step 1
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Step 1: Review Your Priority Actions (Previous Page)', 0, 1)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        step1 = "Focus on the \'START HERE\' priority first. Pick one concrete action from that category and implement it in your next assessment iteration. Do not try to fix everything at once."
        self.multi_cell(0, 6, step1)
        self.ln(4)
        
        # Step 2
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Step 2: Share This Report', 0, 1)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        step2 = "Bring this PDF to your next curriculum planning meeting. Use the reflection questions in each category to start conversations with colleagues about assessment redesign. The evidence quotes help justify changes to leadership."
        self.multi_cell(0, 6, step2)
        self.ln(4)
        
        # Step 3
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Step 3: Re-Audit After Changes', 0, 1)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        step3 = "After implementing improvements, run your revised assessment through the tool again at https://integrity-debt-audit.streamlit.app/ to measure progress. Track your score over time."
        self.multi_cell(0, 6, step3)
        self.ln(4)
        
        # Step 4
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Step 4: Get Support If Needed', 0, 1)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        step4 = "If you are facing institutional resistance or need help prioritising across multiple modules, consider booking a strategy call. I work with departments to implement systematic curriculum redesign."
        self.multi_cell(0, 6, step4)
        self.ln(8)
        
        # Timeline box
        self.set_font('helvetica', 'B', 11)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Suggested Timeline', 0, 1)
        self.ln(2)
        
        self.set_fill_color(*self.bg_cream)
        self.set_draw_color(200, 200, 200)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        timeline = "Week 1: Implement one action from your top priority\nWeek 4: Review impact with students\nWeek 8: Implement second priority action\nEnd of semester: Re-audit and measure improvement"
        self.multi_cell(0, 6, timeline, 1, 'L', True)
        self.ln(8)
    def add_citation_box(self):
        """Academic citation - all in one box"""
        self.check_page_break(30)
        self.set_font('helvetica', 'B', 12)
        self.set_text_color(*self.primary_color)
        self.cell(0, 8, 'Cite This Framework', 0, 1)
        self.ln(2)
        
        # Everything inside one bordered box
        box_y = self.get_y()
        self.set_fill_color(250, 250, 250)
        self.set_draw_color(200, 200, 200)
        
        # Citation text
        self.set_xy(22, box_y + 3)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        self.cell(0, 5, 'Illingworth, S. (2026). The Integrity Debt Audit.', 0, 1)
        
        # URL in blue
        self.set_x(22)
        self.set_font('helvetica', 'U', 10)
        self.set_text_color(0, 0, 255)
        self.cell(0, 5, 'https://integrity-debt-audit.streamlit.app/', 0, 1)
        
        # Draw box around everything
        box_height = self.get_y() - box_y + 3
        self.rect(20, box_y, 170, box_height, 'D')
        self.set_fill_color(250, 250, 250)
        self.rect(20, box_y, 170, box_height, 'F')
        
        # Redraw text on top
        self.set_xy(22, box_y + 3)
        self.set_font('helvetica', '', 10)
        self.set_text_color(*self.text_color_val)
        self.cell(0, 5, 'Illingworth, S. (2026). The Integrity Debt Audit.', 0, 1)
        self.set_x(22)
        self.set_font('helvetica', 'U', 10)
        self.set_text_color(0, 0, 255)
        self.cell(0, 5, 'https://integrity-debt-audit.streamlit.app/', 0, 1)
        
        self.ln(3)

# 3. Utilities
def extract_text(uploaded_file):
    """Extract text from PDF or DOCX files with better error handling"""
    text = ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            for page in reader.pages: 
                text += page.extract_text() or ""
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs: 
                if para.text.strip(): 
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): 
                            text += cell.text + " "
                    text += "\n"
        
        if not text.strip():
            st.error("File appears to be empty or unreadable.")
            return None
            
    except Exception as e: 
        st.error(f"Extraction error: {e}")
        return None
    
    return text

def scrape_url(url):
    """Scrape URL with better security and error handling"""
    try:
        # Add user agent to avoid being blocked
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, timeout=10, headers=headers, allow_redirects=True)
        response.raise_for_status()  # Raise error for bad status codes
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        tags = soup.find_all(['p', 'h1', 'h2', 'h3', 'li', 'td'])
        text = "\n".join([t.get_text().strip() for t in tags if t.get_text().strip()])
        
        if not text:
            return "Error: Could not extract text from URL"
            
        return text
    except Exception as e: 
        return f"Error: {str(e)}"

def clean_json_string(raw_string):
    """Clean JSON response from Gemini with better error handling"""
    # Remove markdown code blocks
    cleaned = re.sub(r'```json\s*|\s*```', '', raw_string)
    # Try to find JSON object
    match = re.search(r'\{.*\}', cleaned, re.DOTALL)
    if match:
        return match.group(0)
    return cleaned.strip()

@st.cache_resource
def discover_model(api_key):
    """Discover available Gemini model with fallback"""
    genai.configure(api_key=api_key)
    try:
        models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        flash = [m for m in models if 'flash' in m.lower()]
        return flash[0] if flash else models[0]
    except:
        return 'models/gemini-1.5-flash'

# CORRECT CATEGORIES FROM THE PDF
INTEGRITY_CATEGORIES = [
    "Final product weighting",
    "Iterative documentation", 
    "Contextual specificity",
    "Reflective criticality",
    "Temporal friction",
    "Multimodal evidence",
    "Explicit AI interrogation",
    "Real-time defence",
    "Social and collaborative labour",
    "Data recency"
]

# Category descriptions for the AI prompt
CATEGORY_DESCRIPTIONS = {
    "Final product weighting": "Does the assessment reward the learning process over the final product? Score 1 (single end-of-term submission) to 5 (multiple formative stages).",
    "Iterative documentation": "Does the assessment require evidence of the messy middle of learning? Score 1 (polished PDF only) to 5 (mandatory brain-dumps, mind maps, rejected ideas).",
    "Contextual specificity": "Is the assessment tied to specific local/classroom contexts that AI cannot access? Score 1 (broad theoretical questions) to 5 (unique in-class discussions).",
    "Reflective criticality": "Does the assessment require deep personal synthesis? Score 1 (generic professional reflection) to 5 (narrative on emotional reactions).",
    "Temporal friction": "Is it physically impossible to complete quickly? Score 1 (can be done in one night) to 5 (longitudinal study over weeks).",
    "Multimodal evidence": "Does the assessment require non-text outputs? Score 1 (standard Word document) to 5 (audio, physical models, hand-drawn).",
    "Explicit AI interrogation": "Does the assessment require students to critique AI outputs? Score 1 (AI ignored or banned) to 5 (generate and critique AI drafts).",
    "Real-time defence": "Does the assessment include live interaction? Score 1 (entirely asynchronous) to 5 (mandatory viva with Q&A).",
    "Social and collaborative labour": "Does the assessment require verified group work? Score 1 (entirely solitary work) to 5 (observed collaboration with peer review).",
    "Data recency": "Does the assessment engage with very recent events/data? Score 1 (static concepts from decades ago) to 5 (last fortnight)."
}

# 4. Interface
st.title("Integrity Debt Diagnostic")
st.caption("🔒 Privacy Statement: This tool is stateless. Assessment briefs are processed in-memory.")

st.markdown("""
### What is this tool?

The **Integrity Debt Audit** helps you identify if your assessments can be easily automated by AI. Many traditional assignments (essays, reports, literature reviews) are now vulnerable to being completed by AI in minutes rather than through genuine student learning. This creates what I call **Integrity Debt**: the gap between what you think you're assessing and what students can now automate.

This diagnostic evaluates your assessment brief across **10 evidence-based categories** that distinguish human learning from AI automation. Each category is scored from 1 (easily automated) to 5 (resilient to AI), giving you a total integrity score out of 50.

### Why use this audit?

- **Stop chasing ghosts with AI detectors** – They don't work reliably, and students know how to evade them
- **Fix the curriculum, not the students** – High scores indicate structural problems with assessment design, not moral failures
- **Protect institutional reputation** – Awarding degrees for AI-generated work threatens the value of your qualifications
- **Design AI-resilient assessments** – Get specific, actionable feedback on how to rebuild assessment integrity

### Important: This is a screening tool

**The browser view provides a preliminary triage only.** For the full diagnostic value, you must download the PDF report, which includes:
- Detailed evidence quotes from your assessment that justify each score
- Pedagogical critiques explaining *why* each category scored as it did  
- Reflective questions to guide curriculum redesign conversations
- Strategic recommendations prioritised by impact

Think of this screen as a medical triage; the PDF is the full diagnostic report you'd discuss with colleagues.

### Need support with implementation?

If your institution needs help interpreting results or redesigning high-stakes assessments, I offer bespoke consultancy.

**Book a strategy call to plan your curriculum redesign:** [sam.illingworth@gmail.com](mailto:sam.illingworth@gmail.com)

Join the Slow AI community for ongoing insights: [theslowai.substack.com](https://theslowai.substack.com)
""")

st.divider()

c1, c2 = st.columns([2, 1])
with c1:
    st.markdown("""
    ### Pre-Audit Checklist
    Include the task description, learning outcomes, and submission formats for accurate results.
    """)
with c2:
    st.info("**The Scoring System**\n* 🟢 4-5: Resilient\n* 🟡 3: Moderate\n* 🔴 1-2: Vulnerable")

st.divider()

# Session State for Modal/Overlay View
if 'audit_complete' not in st.session_state:
    st.session_state.audit_complete = False
if 'user_email' not in st.session_state:
    st.session_state.user_email = ""

with st.container():
    if not st.session_state.audit_complete:
        # Move input type selector OUTSIDE the form so it updates immediately
        st.subheader("2. Assessment Input")
        input_type = st.radio("Choose Input Method:", ["File Upload", "Paste Text or URL"])
        
        with st.form("audit_form"):
            st.subheader("1. Setup")
            sc1, sc2 = st.columns(2)
            with sc1: 
                email_user = st.text_input("Your Email (required):", value=st.session_state.user_email)
            with sc2: 
                expectation = st.selectbox("Predicted Susceptibility:", ["Low", "Medium", "High"])
            
            st.markdown("---")  # Visual separator
            
            uploaded_file = None
            raw_input = ""
            
            if input_type == "File Upload":
                uploaded_file = st.file_uploader("Upload Brief", type=["pdf", "docx"])
            else:  # Paste Text or URL
                raw_input = st.text_area("Paste Assessment Brief or Public URL:", height=200, placeholder="Paste your assessment text here, or enter a URL to a public assessment page...")
            
            submit_button = st.form_submit_button("Generate Diagnostic Report")
    else:
        if st.button("Audit New Brief"):
            st.session_state.audit_complete = False
            st.rerun()

# 5. Logic
api_key = st.secrets.get("GEMINI_API_KEY")

if not st.session_state.audit_complete and 'submit_button' in locals() and submit_button:
    # Validation
    if not api_key:
        st.error("API key configuration missing. Please contact the administrator.")
    elif not email_user or '@' not in email_user:
        st.error("Please provide a valid email address.")
    else:
        # Store email for future audits
        st.session_state.user_email = email_user
        
        # Extract text from input
        final_text = None
        
        if input_type == "File Upload" and uploaded_file:
            final_text = extract_text(uploaded_file)
        elif raw_input:
            if raw_input.startswith("http"):
                with st.spinner("Fetching URL content..."):
                    final_text = scrape_url(raw_input)
                    if final_text.startswith("Error:"):
                        st.error(final_text)
                        final_text = None
            else:
                final_text = raw_input
        
        if not final_text:
            st.error("No assessment content provided or content could not be extracted.")
        elif len(final_text) < 100:
            st.error("Assessment content is too short. Please provide a complete assessment brief.")
        else:
            with st.spinner("Analyzing assessment against Integrity Debt framework..."):
                try:
                    target = discover_model(api_key)
                    model = genai.GenerativeModel(
                        target, 
                        generation_config={"temperature": 0.1, "top_p": 0.2, "top_k": 2}
                    )
                    
                    # Build category descriptions for prompt
                    category_info = "\n".join([f"- {cat}: {desc}" for cat, desc in CATEGORY_DESCRIPTIONS.items()])
                    
                    prompt = f"""You are Professor Sam Illingworth conducting an Integrity Debt Audit for a Higher Education assessment.

CRITICAL INSTRUCTIONS:
1. Analyse the assessment brief against EXACTLY these 10 categories in this exact order:
{category_info}

2. For each category, provide:
   - A score from 1-5 (where 1 = easily automated/vulnerable, 5 = resilient/Slow AI)
   - A critique explaining why you gave this score (keep under 150 words)
   - A dialogue question to help the educator reflect (one sentence)
   - A direct quote from the assessment that supports your score (under 50 words)

3. Your response MUST be valid JSON with this exact structure:
{{
    "doc_context": "Brief title/description of the assessment",
    "top_improvements": ["Improvement 1", "Improvement 2", "Improvement 3"],
    "audit_results": [
        {{
            "category": "Final product weighting",
            "score": 3,
            "critique": "Your analysis here",
            "question": "Reflective question here",
            "quote": "Direct quote from assessment"
        }},
        ... (repeat for all 10 categories in order)
    ]
}}

4. JSON FORMATTING RULES:
   - NO trailing commas before closing braces or brackets
   - Escape ALL quotes within strings using backslash: \\"
   - Keep critique and quote fields SHORT to avoid JSON issues
   - Do NOT include line breaks within string values
   - Use only standard ASCII quotes, not smart quotes

5. IMPORTANT: 
   - Use ONLY the category names listed above
   - Provide exactly 10 results, one for each category
   - Scores must be integers from 1-5
   - Use British English spellings (organise not organize, emphasise not emphasize, etc.)
   - If information is missing for a category, estimate based on typical practices and note this in the critique

Assessment text to analyse (truncated to first 8000 characters):
{final_text[:8000]}

Return ONLY valid JSON with no additional text, markdown formatting, or preamble."""

                    response = model.generate_content(prompt)
                    res_raw = clean_json_string(response.text)
                    
                    try:
                        res_json = json.loads(res_raw)
                    except json.JSONDecodeError as e:
                        # Try to fix common JSON issues
                        try:
                            # Remove trailing commas before closing braces/brackets
                            fixed_json = re.sub(r',\s*}', '}', res_raw)
                            fixed_json = re.sub(r',\s*]', ']', fixed_json)
                            # Escape unescaped quotes in strings
                            res_json = json.loads(fixed_json)
                        except json.JSONDecodeError:
                            st.error(f"The AI response could not be parsed. This sometimes happens with very long assessments. Please try:")
                            st.markdown("""
                            - Shortening your assessment text slightly
                            - Removing any special characters or unusual formatting
                            - Trying again (the AI may succeed on second attempt)
                            """)
                            with st.expander("Show AI Response (for debugging)"):
                                st.code(res_raw)
                            st.stop()
                    
                    # Validate response structure
                    if not isinstance(res_json, dict):
                        st.error("AI returned invalid response format.")
                        st.stop()
                    
                    if "status" in res_json and res_json["status"] == "error":
                        st.error("No clear assessment task could be identified in the provided text.")
                        st.stop()
                    
                    # Store in session state
                    st.session_state.res_json = res_json
                    st.session_state.audit_complete = True
                    st.rerun()
                    
                except Exception as e: 
                    st.error(f"Audit failed: {e}")
                    st.error("This may be due to API limits or connectivity issues. Please try again.")

# Display results
if st.session_state.audit_complete:
    res_json = st.session_state.res_json
    
    # Process audit results with proper validation
    total_score = 0
    final_audit_results = {}
    audit_list = res_json.get("audit_results", [])
    
    # Ensure audit_list is actually a list
    if not isinstance(audit_list, list):
        if isinstance(audit_list, dict):
            audit_list = list(audit_list.values())
        else:
            audit_list = []
    
    # Match results to categories
    for cat_name in INTEGRITY_CATEGORIES:
        # Try to find a match for this category
        match = None
        for item in audit_list:
            if not isinstance(item, dict):
                continue
            item_cat = item.get('category', '').lower()
            if cat_name.lower() in item_cat or item_cat in cat_name.lower():
                match = item
                break
        
        if match:
            # Extract score - try multiple possible fields
            score = 0
            for field in ['score', 'points', 'rating']:
                if field in match:
                    try:
                        score = int(match[field])
                        break
                    except (ValueError, TypeError):
                        # Try to extract number from string
                        score_str = str(match[field])
                        score_match = re.search(r'\d+', score_str)
                        if score_match:
                            score = int(score_match.group(0))
                            break
            
            # Clamp score to valid range
            score = max(1, min(5, score))
            
            final_audit_results[cat_name] = {
                'verified_score': score,
                'critique': match.get('critique', 'No critique provided'),
                'question': match.get('question', 'No question provided'),
                'quote': match.get('quote', 'No quote provided')
            }
            total_score += score
        else:
            # No match found - create placeholder
            final_audit_results[cat_name] = {
                'verified_score': 3,  # Default to middle score
                'critique': 'Insufficient information provided to evaluate this category.',
                'question': 'How could this category be better addressed in your assessment?',
                'quote': 'N/A'
            }
            total_score += 3
    
    # Get context and improvements
    ctx = res_json.get("doc_context", "Assessment Audit")
    imps = res_json.get("top_improvements", ["Review individual categories for specific improvements"])
    
    # Ensure improvements is a list
    if not isinstance(imps, list):
        imps = [str(imps)]
    
    # Calculate susceptibility level - HIGH scores (40-50) are GOOD
    if total_score >= 40:
        susceptibility = "Low (Pedagogical Sovereignty)"
    elif total_score >= 25:
        susceptibility = "Medium (Structural Drift)"
    else:
        susceptibility = "High (Critical Integrity Failure)"
    
    # Display summary
    st.markdown("""
        <style>
        /* Make results section text visible */
        h2, h3 {
            color: #2D4D4A !important;
        }
        [data-testid="stMetricValue"] {
            color: #2D4D4A !important;
            font-size: 2rem !important;
            font-weight: bold !important;
        }
        [data-testid="stMetricLabel"] {
            color: #2D4D4A !important;
            font-weight: 600 !important;
        }
        [data-testid="stExpander"] {
            background-color: #f8f9fa !important;
            border: 1px solid #dee2e6 !important;
        }
        [data-testid="stExpander"] p {
            color: #212529 !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.subheader(f"Diagnostic Focus: {ctx}")
    
    col1, col2 = st.columns(2)
    with col1:
        # HIGH total score (40-50) = GOOD (green), LOW (10-20) = BAD (red)
        score_color = "🟢" if total_score >= 40 else "🟡" if total_score >= 25 else "🔴"
        st.metric("Integrity Score", f"{score_color} {total_score}/50")
    with col2:
        st.metric("AI Susceptibility", susceptibility.split('(')[0].strip())
    
    st.markdown("""
        <p style='color: #856404; background-color: #fff3cd; padding: 12px; border-radius: 4px; border-left: 4px solid #ffc107;'>
        ⚠️ <strong>Note:</strong> Formal PDF contains full detailed report; this view is a summary.
        </p>
    """, unsafe_allow_html=True)
    
    # Generate PDF
    pdf = IntegrityPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.add_summary(susceptibility.split('(')[0].strip(), total_score, imps, ctx)
    
    for cat_name, data in final_audit_results.items():
        score = data['verified_score']
        
        # Get pedagogical context and improvement actions
        pedagogical_context = PEDAGOGICAL_CONTEXT[cat_name]
        actions = IMPROVEMENT_ACTIONS[cat_name]
        
        # Use enhanced method with theory and actions
        pdf.add_category_deep_dive(
            cat_name, 
            score,
            data['critique'], 
            data['question'], 
            data['quote'],
            pedagogical_context,
            actions
        )
    
    # Add score interpretation with visual scale
    pdf.add_score_interpretation(total_score)
    
    # Add action plan page
    pdf.add_action_plan(final_audit_results)
    
    # Add next steps guidance
    pdf.add_next_steps()
    
    # Add citation
    pdf.add_citation_box()
    
    # Add contact box last
    pdf.add_contact_box()
    
    # Download button with custom styling - simple white/black
    st.markdown("""
        <style>
        .stDownloadButton > button {
            background-color: white !important;
            color: #212529 !important;
            border: 2px solid #212529 !important;
            padding: 16px 32px !important;
            font-size: 18px !important;
            font-weight: bold !important;
            border-radius: 8px !important;
            width: 100% !important;
            margin: 20px 0 !important;
        }
        .stDownloadButton > button:hover {
            background-color: #f8f9fa !important;
            border: 2px solid #000000 !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.download_button(
        "📥 Download Full Evidence Report (PDF)", 
        data=bytes(pdf.output()), 
        file_name="Integrity_Debt_Audit.pdf",
        mime="application/pdf"
    )
    
    st.divider()
    
    # Display category breakdown
    st.subheader("Category Breakdown")
    
    for cat_name, data in final_audit_results.items():
        score = data['verified_score']
        # HIGH scores (4-5) are GOOD (resilient) = GREEN
        # LOW scores (1-2) are BAD (vulnerable) = RED
        label = "🟢" if score >= 4 else "🟡" if score == 3 else "🔴"
        
        with st.expander(f"{label} **{cat_name}** ({score}/5)"):
            st.markdown(f"**Critique:** {data['critique']}")
            st.markdown(f"**Question for reflection:** {data['question']}")
            if data['quote'] != 'N/A':
                st.markdown(f"**Evidence:** _{data['quote']}_")
    
    st.divider()
    st.caption("End of summary. For full pedagogical rationale and evidence quotes, please download the PDF report.")
